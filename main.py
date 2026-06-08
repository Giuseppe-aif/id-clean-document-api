from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Header
from typing import Optional
from pathlib import Path
import base64
import os
import re
import tempfile
import io
import json

import cv2
import numpy as np
from PIL import Image, ImageFilter, ExifTags
from rembg import remove, new_session

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.utils import ImageReader


app = FastAPI()

API_KEY = os.getenv("API_KEY", "")

# ── Physical document dimensions ──────────────────────────────────────────────

ID_CARD_WIDTH_MM  = 85.60
ID_CARD_HEIGHT_MM = 53.98

PASSPORT_SCAN_WIDTH_MM  = 176.0
PASSPORT_SCAN_HEIGHT_MM = 125.0

PASSPORT_WIDTH_MM  = PASSPORT_SCAN_HEIGHT_MM   # 125
PASSPORT_HEIGHT_MM = PASSPORT_SCAN_WIDTH_MM    # 176

DPI = 300

# ── Page / layout constants ───────────────────────────────────────────────────

DOCX_PAGE_LEFT_MARGIN_MM   = 20.0
DOCX_PAGE_RIGHT_MARGIN_MM  = 20.0
DOCX_PAGE_TOP_MARGIN_MM    = 20.0
DOCX_PAGE_BOTTOM_MARGIN_MM = 20.0

PAGE_IMAGE_LEFT_MARGIN_MM = 40.0
PDF_CARD_MARGIN_MM        = 8.0

# ── Shadow settings ───────────────────────────────────────────────────────────

SHADOW_OFFSET_MM = 2.0
SHADOW_BLUR_MM   = 3.0
SHADOW_OPACITY   = 160
SHADOW_COLOR     = (80, 80, 80)

# ── rembg tunables ────────────────────────────────────────────────────────────

ALPHA_THRESHOLD       = 10
TIGHT_CROP_PADDING_PX = 4


# ── rembg session ─────────────────────────────────────────────────────────────

_REMBG_SESSION = None

def get_rembg_session():
    global _REMBG_SESSION
    if _REMBG_SESSION is None:
        _REMBG_SESSION = new_session("isnet-general-use")
    return _REMBG_SESSION


# ── Health endpoints ──────────────────────────────────────────────────────────

@app.get("/")
def home():
    return {"status": "ok", "message": "ID Clean Document API is running"}

@app.get("/health")
def health():
    return {"status": "healthy"}


# ── General utilities ─────────────────────────────────────────────────────────

def safe_file_part(value: str) -> str:
    value = str(value or "").strip()
    value = re.sub(r"[^A-Za-z0-9_-]+", "_", value)
    value = value.strip("_")
    return value or "clean_document"

def mm_to_px(mm_value: float, dpi: int = DPI) -> int:
    return int(round(mm_value / 25.4 * dpi))


# ── EXIF-aware image reading ──────────────────────────────────────────────────

def read_image_exif_aware(path: Path) -> np.ndarray:
    """
    Read image respecting EXIF orientation.
    Returns BGR ndarray (OpenCV convention).
    """
    pil_img = Image.open(path)
    try:
        exif = pil_img._getexif()
        if exif:
            orientation_key = next(
                (k for k, v in ExifTags.TAGS.items() if v == "Orientation"), None
            )
            if orientation_key and orientation_key in exif:
                orientation = exif[orientation_key]
                rotations = {3: 180, 6: -90, 8: 90}
                if orientation in rotations:
                    pil_img = pil_img.rotate(rotations[orientation], expand=True)
                elif orientation in {2, 4, 5, 7}:
                    pil_img = pil_img.transpose(Image.FLIP_LEFT_RIGHT)
                    if orientation in {5, 7}:
                        pil_img = pil_img.rotate(90 if orientation == 5 else -90, expand=True)
    except Exception:
        pass
    pil_img = pil_img.convert("RGB")
    return cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)


def bgr_to_pil(image_bgr: np.ndarray) -> Image.Image:
    return Image.fromarray(cv2.cvtColor(image_bgr, cv2.COLOR_BGR2RGB))


# ── Perspective correction using Claude corners ───────────────────────────────

def parse_corners(corners_json: Optional[str]) -> Optional[np.ndarray]:
    """
    Parse corner coordinates sent from n8n (Claude Vision output).
    Expects JSON: {"top_left": [x,y], "top_right": [x,y],
                   "bottom_right": [x,y], "bottom_left": [x,y],
                   "rotation_needed": 0}
    Returns float32 (4,2) array or None.
    """
    if not corners_json:
        return None
    try:
        data = json.loads(corners_json)
        pts = np.array([
            data["top_left"],
            data["top_right"],
            data["bottom_right"],
            data["bottom_left"],
        ], dtype="float32")
        return pts
    except Exception:
        return None

def parse_rotation(corners_json: Optional[str]) -> int:
    """Extract rotation_needed from corners JSON. Returns 0 if missing."""
    if not corners_json:
        return 0
    try:
        data = json.loads(corners_json)
        return int(data.get("rotation_needed", 0))
    except Exception:
        return 0

def order_points(pts: np.ndarray) -> np.ndarray:
    rect = np.zeros((4, 2), dtype="float32")
    s    = pts.sum(axis=1)
    rect[0] = pts[np.argmin(s)]
    rect[2] = pts[np.argmax(s)]
    diff    = np.diff(pts, axis=1)
    rect[1] = pts[np.argmin(diff)]
    rect[3] = pts[np.argmax(diff)]
    return rect

def four_point_transform(image: np.ndarray, pts: np.ndarray) -> np.ndarray:
    """Warp image to flat rectangle using 4 corner points."""
    rect       = order_points(pts)
    tl, tr, br, bl = rect
    max_width  = int(max(np.linalg.norm(br - bl), np.linalg.norm(tr - tl)))
    max_height = int(max(np.linalg.norm(tr - br), np.linalg.norm(tl - bl)))
    if max_width < 50 or max_height < 50:
        return image
    dst = np.array([
        [0,             0             ],
        [max_width - 1, 0             ],
        [max_width - 1, max_height - 1],
        [0,             max_height - 1],
    ], dtype="float32")
    matrix = cv2.getPerspectiveTransform(rect, dst)
    return cv2.warpPerspective(
        image, matrix, (max_width, max_height),
        borderValue=(255, 255, 255),
    )

def apply_rotation(image: np.ndarray, rotation_needed: int) -> np.ndarray:
    """Apply clockwise rotation as specified by Claude."""
    if rotation_needed == 90:
        return cv2.rotate(image, cv2.ROTATE_90_CLOCKWISE)
    elif rotation_needed == -90 or rotation_needed == 270:
        return cv2.rotate(image, cv2.ROTATE_90_COUNTERCLOCKWISE)
    elif rotation_needed == 180:
        return cv2.rotate(image, cv2.ROTATE_180)
    return image


# ── Background removal ────────────────────────────────────────────────────────

def remove_background(image_bgr: np.ndarray) -> Image.Image:
    """
    Use rembg (ISNet) to remove background after perspective correction.
    At this point the card is already a clean rectangle so rembg only
    needs to separate card from any residual border — much easier task.
    Returns RGBA PIL image.
    """
    pil_rgb  = bgr_to_pil(image_bgr)
    rgba_out = remove(pil_rgb, session=get_rembg_session())

    # Hard rectangular clip from alpha bounding box — sharp corners
    alpha  = np.array(rgba_out.split()[3])
    coords = np.argwhere(alpha > ALPHA_THRESHOLD)

    if coords.size == 0:
        return rgba_out

    y0, x0 = coords.min(axis=0)
    y1, x1 = coords.max(axis=0)
    h, w   = alpha.shape

    y0 = max(y0 - TIGHT_CROP_PADDING_PX, 0)
    x0 = max(x0 - TIGHT_CROP_PADDING_PX, 0)
    y1 = min(y1 + TIGHT_CROP_PADDING_PX, h - 1)
    x1 = min(x1 + TIGHT_CROP_PADDING_PX, w - 1)

    rgba_cropped = rgba_out.crop((x0, y0, x1 + 1, y1 + 1))

    # Binarise alpha — hard edges, no feathering
    r, g, b, a = rgba_cropped.split()
    a_arr = np.array(a)
    a_arr[a_arr > ALPHA_THRESHOLD] = 255
    a_arr[a_arr <= ALPHA_THRESHOLD] = 0
    rgba_hard = Image.merge("RGBA", (r, g, b, Image.fromarray(a_arr)))

    return rgba_hard


# ── Orientation helpers ───────────────────────────────────────────────────────

def force_landscape(rgba: Image.Image) -> Image.Image:
    w, h = rgba.size
    if h > w:
        return rgba.rotate(-90, expand=True)
    return rgba

def force_passport_orientation(rgba: Image.Image) -> Image.Image:
    w, h = rgba.size
    if h > w:
        rgba = rgba.rotate(90, expand=True)
    return rgba.rotate(90, expand=True)


# ── Exact-size renderer ───────────────────────────────────────────────────────

def render_card_exact(
    rgba: Image.Image,
    target_width_mm: float,
    target_height_mm: float,
    doc_type: str = "id",
) -> Image.Image:
    if doc_type == "passport":
        rgba = force_passport_orientation(rgba)
    else:
        rgba = force_landscape(rgba)

    target_w_px = mm_to_px(target_width_mm)
    target_h_px = mm_to_px(target_height_mm)
    return rgba.resize((target_w_px, target_h_px), Image.LANCZOS)


# ── Drop shadow compositor ────────────────────────────────────────────────────

def add_drop_shadow(card_rgba: Image.Image) -> Image.Image:
    """
    Composite RGBA card onto white canvas with Gaussian drop shadow.
    Shadow follows the card's alpha so no white-box halo is visible.
    """
    cw, ch = card_rgba.size

    shadow_offset_px = mm_to_px(SHADOW_OFFSET_MM)
    shadow_blur_px   = mm_to_px(SHADOW_BLUR_MM)

    pad      = shadow_blur_px * 2 + shadow_offset_px
    canvas_w = cw + pad
    canvas_h = ch + pad
    card_x   = pad // 2 - shadow_offset_px // 2
    card_y   = pad // 2 - shadow_offset_px // 2
    shadow_x = card_x + shadow_offset_px
    shadow_y = card_y + shadow_offset_px

    base = Image.new("RGBA", (canvas_w, canvas_h), (255, 255, 255, 255))

    shadow_layer = Image.new("RGBA", (canvas_w, canvas_h), (255, 255, 255, 0))
    shadow_card  = Image.new("RGBA", (cw, ch),
                             (SHADOW_COLOR[0], SHADOW_COLOR[1], SHADOW_COLOR[2], 0))
    shadow_alpha = card_rgba.split()[3].point(lambda x: int(x * SHADOW_OPACITY / 255))
    shadow_card.putalpha(shadow_alpha)
    shadow_layer.paste(shadow_card, (shadow_x, shadow_y))
    shadow_layer = shadow_layer.filter(ImageFilter.GaussianBlur(radius=shadow_blur_px))

    base = Image.alpha_composite(base, shadow_layer)
    base.paste(card_rgba, (card_x, card_y), mask=card_rgba.split()[3])

    return base.convert("RGB")


# ── PIL → ReportLab ImageReader ───────────────────────────────────────────────

def pil_to_image_reader(pil_img: Image.Image) -> ImageReader:
    buf = io.BytesIO()
    pil_img.save(buf, format="PNG", dpi=(DPI, DPI))
    buf.seek(0)
    return ImageReader(buf)


# ── Full processing pipeline ──────────────────────────────────────────────────

def process_image(
    input_path: Path,
    output_path: Path,
    target_width_mm: float,
    target_height_mm: float,
    doc_type: str = "id",
    corners_json: Optional[str] = None,
) -> dict:
    """
    Pipeline:
    1. EXIF-aware read
    2. Perspective correction using Claude's corner coordinates
    3. Rotation correction using Claude's rotation_needed value
    4. rembg background removal (easier task post-correction)
    5. Hard rectangular clip — sharp corners, transparent background
    6. Exact physical resize
    7. Drop shadow composite
    8. Save PNG at 300 DPI
    """
    image = read_image_exif_aware(input_path)

    # Step 2: perspective correction from Claude corners
    pts = parse_corners(corners_json)
    if pts is not None:
        image  = four_point_transform(image, pts)
        method = "claude_corners_perspective_corrected"
    else:
        method = "no_corners_fallback"

    # Step 3: rotation correction
    rotation = parse_rotation(corners_json)
    if rotation != 0:
        image  = apply_rotation(image, rotation)
        method += f"_rotated_{rotation}"

    # Step 4-5: background removal + hard clip
    card_rgba = remove_background(image)

    # Step 6: exact physical size + orientation
    card_exact = render_card_exact(card_rgba, target_width_mm, target_height_mm, doc_type=doc_type)

    # Step 7-8: shadow + save
    final_pil = add_drop_shadow(card_exact)
    final_pil.save(output_path, dpi=(DPI, DPI))

    return {
        "input_file":  input_path.name,
        "output_file": output_path.name,
        "method":      method,
    }


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _image_display_width_mm(base_mm: float) -> float:
    pad_px = mm_to_px(SHADOW_BLUR_MM) * 2 + mm_to_px(SHADOW_OFFSET_MM)
    pad_mm = pad_px / (DPI / 25.4)
    return base_mm + pad_mm


def add_image_paragraph(doc: Document, image_path: Path, width_mm: float):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    extra_indent_mm = max(PAGE_IMAGE_LEFT_MARGIN_MM - DOCX_PAGE_LEFT_MARGIN_MM, 0)
    paragraph.paragraph_format.left_indent = Cm(extra_indent_mm / 10.0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_mm / 10.0))


def create_docx(
    docx_path: Path,
    doc_type: str,
    front_image: Path,
    back_image: Optional[Path],
):
    doc     = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(DOCX_PAGE_TOP_MARGIN_MM    / 10.0)
    section.bottom_margin = Cm(DOCX_PAGE_BOTTOM_MARGIN_MM / 10.0)
    section.left_margin   = Cm(DOCX_PAGE_LEFT_MARGIN_MM   / 10.0)
    section.right_margin  = Cm(DOCX_PAGE_RIGHT_MARGIN_MM  / 10.0)

    if doc_type == "id":
        w_mm = _image_display_width_mm(ID_CARD_WIDTH_MM)
        add_image_paragraph(doc, front_image, w_mm)
        doc.add_paragraph("")
        add_image_paragraph(doc, back_image, w_mm)
    elif doc_type == "passport":
        w_mm = _image_display_width_mm(PASSPORT_WIDTH_MM)
        add_image_paragraph(doc, front_image, w_mm)
    else:
        raise ValueError(f"Unsupported document type: {doc_type}")

    doc.save(docx_path)


# ── PDF helpers ───────────────────────────────────────────────────────────────

def create_pdf(
    pdf_path: Path,
    doc_type: str,
    front_image: Path,
    back_image: Optional[Path],
):
    page_w, page_h = A4
    c = rl_canvas.Canvas(str(pdf_path), pagesize=A4)
    x = PAGE_IMAGE_LEFT_MARGIN_MM * mm

    def draw(img_path: Path, y_top_mm: float) -> float:
        pil_img = Image.open(img_path)
        w_mm    = pil_img.width  / (DPI / 25.4)
        h_mm    = pil_img.height / (DPI / 25.4)
        y_pt    = page_h - y_top_mm * mm - h_mm * mm
        reader  = pil_to_image_reader(pil_img)
        c.drawImage(reader, x, y_pt, width=w_mm * mm, height=h_mm * mm, mask="auto")
        return h_mm

    if doc_type == "id":
        y = DOCX_PAGE_TOP_MARGIN_MM + PDF_CARD_MARGIN_MM
        h = draw(front_image, y)
        draw(back_image, y + h + PDF_CARD_MARGIN_MM)
    elif doc_type == "passport":
        draw(front_image, DOCX_PAGE_TOP_MARGIN_MM + PDF_CARD_MARGIN_MM)
    else:
        raise ValueError(f"Unsupported document type: {doc_type}")

    c.showPage()
    c.save()


# ── Upload / base64 utils ─────────────────────────────────────────────────────

async def save_upload(upload_file: UploadFile, destination: Path):
    destination.write_bytes(await upload_file.read())

def file_to_base64(path: Path) -> str:
    return base64.b64encode(path.read_bytes()).decode("utf-8")


# ── Endpoints ─────────────────────────────────────────────────────────────────

@app.post("/process-document")
async def process_document(
    first_name:       str                  = Form(...),
    last_name:        str                  = Form(...),
    doc_type:         str                  = Form(...),
    output_base_name: str                  = Form(...),
    front_image:      UploadFile           = File(...),
    back_image:       Optional[UploadFile] = File(None),
    front_corners:    Optional[str]        = Form(None),
    back_corners:     Optional[str]        = Form(None),
    x_api_key:        Optional[str]        = Header(None),
):
    if API_KEY and x_api_key != API_KEY:
        raise HTTPException(status_code=401, detail="Invalid API key")

    doc_type = str(doc_type or "").lower().strip()
    if doc_type not in ["id", "passport"]:
        raise HTTPException(status_code=400, detail="doc_type must be 'id' or 'passport'")
    if doc_type == "id" and back_image is None:
        raise HTTPException(status_code=400, detail="Back image is required for ID documents")

    output_base_name = safe_file_part(output_base_name)

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp         = Path(tmpdir)
        front_input = tmp / "front_input"
        back_input  = tmp / "back_input"

        await save_upload(front_image, front_input)
        if back_image:
            await save_upload(back_image, back_input)

        processed_info = []

        if doc_type == "id":
            front_processed = tmp / "front_processed.png"
            back_processed  = tmp / "back_processed.png"
            processed_info.append(
                process_image(
                    front_input, front_processed,
                    ID_CARD_WIDTH_MM, ID_CARD_HEIGHT_MM,
                    doc_type="id",
                    corners_json=front_corners,
                )
            )
            processed_info.append(
                process_image(
                    back_input, back_processed,
                    ID_CARD_WIDTH_MM, ID_CARD_HEIGHT_MM,
                    doc_type="id",
                    corners_json=back_corners,
                )
            )
        else:
            front_processed = tmp / "passport_processed.png"
            back_processed  = None
            processed_info.append(
                process_image(
                    front_input, front_processed,
                    PASSPORT_WIDTH_MM, PASSPORT_HEIGHT_MM,
                    doc_type="passport",
                    corners_json=front_corners,
                )
            )

        docx_filename = f"{output_base_name}.docx"
        pdf_filename  = f"{output_base_name}.pdf"
        docx_path     = tmp / docx_filename
        pdf_path      = tmp / pdf_filename

        create_docx(
            docx_path=docx_path, doc_type=doc_type,
            front_image=front_processed, back_image=back_processed,
        )
        create_pdf(
            pdf_path=pdf_path, doc_type=doc_type,
            front_image=front_processed, back_image=back_processed,
        )

        return {
            "status":           "success",
            "doc_type":         doc_type,
            "first_name":       first_name,
            "last_name":        last_name,
            "output_base_name": output_base_name,
            "docx_filename":    docx_filename,
            "pdf_filename":     pdf_filename,
            "docx_base64":      file_to_base64(docx_path),
            "pdf_base64":       file_to_base64(pdf_path),
            "processed_images": processed_info,
        }


@app.post("/process-document-test")
async def process_document_test(
    first_name:       str                  = Form(...),
    last_name:        str                  = Form(...),
    doc_type:         str                  = Form(...),
    output_base_name: str                  = Form(...),
    front_image:      UploadFile           = File(...),
    back_image:       Optional[UploadFile] = File(None),
    front_corners:    Optional[str]        = Form(None),
    back_corners:     Optional[str]        = Form(None),
    x_api_key:        Optional[str]        = Header(None),
):
    if API_KEY and x_api_key != API_KEY:
        raise HTTPException(status_code=401, detail="Invalid API key")
    if doc_type == "id" and back_image is None:
        raise HTTPException(status_code=400, detail="Back image is required for ID documents")

    return {
        "status":           "received",
        "first_name":       first_name,
        "last_name":        last_name,
        "doc_type":         doc_type,
        "output_base_name": output_base_name,
        "front_filename":   front_image.filename,
        "back_filename":    back_image.filename if back_image else None,
        "front_corners":    front_corners,
        "back_corners":     back_corners,
    }


VERSION = "2025-06-08-v9"

@app.get("/version")
def version():
    return {"version": VERSION}
